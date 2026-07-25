"""Product management API routes — AI-enhanced product creation and management."""

import io
import json
import logging
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from typing import Optional

logger = logging.getLogger(__name__)

product_router = APIRouter(prefix="/agent", tags=["product"])


class EnhanceProductRequest(BaseModel):
    product_id: str
    name: str
    description: str = ""
    short_description: str = ""
    category_name: str = ""
    models: str = ""
    file_descriptions: str = ""


class EnhanceProductResponse(BaseModel):
    success: bool
    product_id: str
    enhanced: Optional[dict] = None
    message: str = ""
    error: Optional[str] = None


@product_router.post("/enhance-product", response_model=EnhanceProductResponse)
async def enhance_product(body: EnhanceProductRequest):
    """AI-enhanced product detail generation.

    Takes raw product data + file descriptions and uses LLM to generate
    professional descriptions, specs, and SEO metadata.
    """
    try:
        from app.tools.product_enhance import enhance_product_details

        result = await enhance_product_details.ainvoke({
            "product_id": body.product_id,
            "name": body.name,
            "short_description": body.short_description,
            "description": body.description,
            "category_name": body.category_name,
            "models": body.models,
            "file_descriptions": body.file_descriptions,
        })

        data = json.loads(result)
        if data.get("error"):
            return EnhanceProductResponse(
                success=False,
                product_id=body.product_id,
                message=data["error"],
                error=data["error"],
            )

        return EnhanceProductResponse(
            success=True,
            product_id=body.product_id,
            enhanced=data.get("enhanced"),
            message=data.get("message", "Product enhanced successfully."),
        )

    except Exception as e:
        logger.error(f"Product enhancement failed: {e}", exc_info=True)
        return EnhanceProductResponse(
            success=False,
            product_id=body.product_id,
            message=f"Enhancement failed: {str(e)[:200]}",
            error=str(e),
        )


class CreateCategoryRequest(BaseModel):
    name: str
    slug: str = ""


class CreateCategoryResponse(BaseModel):
    success: bool
    category_id: str = ""
    name: str = ""
    slug: str = ""
    message: str = ""


@product_router.post("/create-category", response_model=CreateCategoryResponse)
async def create_category_endpoint(body: CreateCategoryRequest):
    """Create a new product category."""
    try:
        from app.tools.product_enhance import create_category

        result = await create_category.ainvoke({
            "name": body.name,
            "slug": body.slug,
        })

        data = json.loads(result)
        return CreateCategoryResponse(
            success=data.get("success", False),
            category_id=data.get("category_id", ""),
            name=data.get("name", ""),
            slug=data.get("slug", ""),
            message=data.get("message", ""),
        )

    except Exception as e:
        logger.error(f"Category creation failed: {e}", exc_info=True)
        return CreateCategoryResponse(
            success=False,
            message=f"Category creation failed: {str(e)[:200]}",
        )


@product_router.get("/list-categories")
async def list_categories_endpoint():
    """List all active product categories."""
    try:
        from app.tools.product_enhance import list_categories

        result = await list_categories.ainvoke({})
        return {"success": True, "data": result}

    except Exception as e:
        logger.error(f"Failed to list categories: {e}")
        return {"success": False, "error": str(e)}


# ── Product Info Extraction from PDF/DOCX ───────────────


class ExtractProductInfoRequest(BaseModel):
    product_id: str
    file_text: str
    file_name: str = ""


class ExtractProductInfoResponse(BaseModel):
    success: bool
    product_id: str = ""
    extracted: Optional[dict] = None
    message: str = ""
    error: Optional[str] = None


@product_router.post("/extract-product-info", response_model=ExtractProductInfoResponse)
async def extract_product_info_endpoint(body: ExtractProductInfoRequest):
    """Extract product information from PDF/DOCX text using LLM.

    Used by both the chat agent and the Add Product modal step 3.
    Stores extracted data in llmExtractedData for admin review.
    """
    try:
        from app.tools.product_admin import extract_product_info

        result = await extract_product_info.ainvoke({
            "product_id": body.product_id,
            "file_text": body.file_text,
            "file_name": body.file_name,
        })

        data = json.loads(result)
        if data.get("error"):
            return ExtractProductInfoResponse(
                success=False,
                product_id=body.product_id,
                message=data["error"],
                error=data["error"],
            )

        return ExtractProductInfoResponse(
            success=True,
            product_id=body.product_id,
            extracted=data.get("extracted"),
            message=data.get("message", "Product info extracted."),
        )

    except Exception as e:
        logger.error(f"Product info extraction failed: {e}", exc_info=True)
        return ExtractProductInfoResponse(
            success=False,
            product_id=body.product_id,
            message=f"Extraction failed: {str(e)[:200]}",
            error=str(e),
        )


# ── Product Upload Media ─────────────────────────────────


class UploadProductMediaRequest(BaseModel):
    product_id: str
    file_name: str
    file_bytes_b64: str
    media_type: str = "image"
    alt: str = ""


class UploadProductMediaResponse(BaseModel):
    success: bool
    product_id: str = ""
    url: str = ""
    message: str = ""
    error: Optional[str] = None


@product_router.post("/upload-product-media", response_model=UploadProductMediaResponse)
async def upload_product_media_endpoint(body: UploadProductMediaRequest):
    """Upload media (image/video) to S3/R2 and attach to a product."""
    try:
        from app.tools.product_admin import upload_product_media

        result = await upload_product_media.ainvoke({
            "product_id": body.product_id,
            "file_name": body.file_name,
            "file_bytes_b64": body.file_bytes_b64,
            "media_type": body.media_type,
            "alt": body.alt,
        })

        data = json.loads(result)
        if data.get("error"):
            return UploadProductMediaResponse(
                success=False,
                product_id=body.product_id,
                message=data["error"],
                error=data["error"],
            )

        return UploadProductMediaResponse(
            success=True,
            product_id=body.product_id,
            url=data.get("url", ""),
            message=data.get("message", "Media uploaded."),
        )

    except Exception as e:
        logger.error(f"Product media upload failed: {e}", exc_info=True)
        return UploadProductMediaResponse(
            success=False,
            product_id=body.product_id,
            message=f"Upload failed: {str(e)[:200]}",
            error=str(e),
        )


# ── Extract from Uploaded File (combined: upload + extract) ────────


class ExtractFromFileResponse(BaseModel):
    success: bool
    product_id: str = ""
    extracted: Optional[dict] = None
    files_processed: int = 0
    message: str = ""
    error: Optional[str] = None


@product_router.post("/extract-from-file", response_model=ExtractFromFileResponse)
async def extract_from_file_endpoint(
    product_id: str,
    file_name: str = "",
    file_text: str = "",
):
    """Extract product information from an uploaded PDF/DOCX file.

    Accepts pre-extracted text (from frontend PDF.js or backend PyPDF2)
    and calls extract_product_info to produce structured data.

    Used by the Add Product modal step 3.
    """
    try:
        from app.tools.product_admin import extract_product_info

        if not file_text or not file_text.strip():
            return ExtractFromFileResponse(
                success=False,
                product_id=product_id,
                message="No text content provided for extraction.",
                error="Empty file text",
            )

        result = await extract_product_info.ainvoke({
            "product_id": product_id,
            "file_text": file_text,
            "file_name": file_name,
        })

        data = json.loads(result)
        if data.get("error"):
            return ExtractFromFileResponse(
                success=False,
                product_id=product_id,
                message=data["error"],
                error=data["error"],
            )

        return ExtractFromFileResponse(
            success=True,
            product_id=product_id,
            extracted=data.get("extracted"),
            files_processed=1,
            message=data.get("message", "File processed and extraction complete."),
        )

    except Exception as e:
        logger.error(f"Extract from file failed: {e}", exc_info=True)
        return ExtractFromFileResponse(
            success=False,
            product_id=product_id,
            message=f"Extraction failed: {str(e)[:200]}",
            error=str(e),
        )


# ── Extract from File Upload (multipart: handles PDF/DOCX) ──────────


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2."""
    try:
        from PyPDF2 import PdfReader

        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


def _extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text from a file based on its extension."""
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return _extract_text_from_docx(file_bytes)
    elif lower_name.endswith(".txt") or lower_name.endswith(".csv") or lower_name.endswith(".md"):
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return ""
    return ""


@product_router.post("/extract-from-upload", response_model=ExtractFromFileResponse)
async def extract_from_upload_endpoint(
    product_id: str = Form(...),
    file: UploadFile = File(...),
):
    """Upload a PDF/DOCX/TXT file, extract text, and run LLM extraction.

    Used by the Add Product modal step 3 — accepts the actual file upload,
    extracts text server-side, then calls extract_product_info.
    """
    try:
        from app.tools.product_admin import extract_product_info

        file_bytes = await file.read()

        # Extract text based on file type
        file_text = _extract_text_from_file(file_bytes, file.filename or "")

        if not file_text or not file_text.strip():
            return ExtractFromFileResponse(
                success=False,
                product_id=product_id,
                message=f"Could not extract text from '{file.filename}'. The file may be image-based or empty.",
                error="No extractable text",
            )

        # Truncate very long texts
        if len(file_text) > 20000:
            file_text = file_text[:20000] + "\n\n[Text truncated at 20000 chars]"

        result = await extract_product_info.ainvoke({
            "product_id": product_id,
            "file_text": file_text,
            "file_name": file.filename or "",
        })

        data = json.loads(result)
        if data.get("error"):
            return ExtractFromFileResponse(
                success=False,
                product_id=product_id,
                message=data["error"],
                error=data["error"],
            )

        return ExtractFromFileResponse(
            success=True,
            product_id=product_id,
            extracted=data.get("extracted"),
            files_processed=1,
            message=data.get("message", "File processed successfully."),
        )

    except Exception as e:
        logger.error(f"Extract from upload failed: {e}", exc_info=True)
        return ExtractFromFileResponse(
            success=False,
            product_id=product_id,
            message=f"Extraction failed: {str(e)[:200]}",
            error=str(e),
        )
